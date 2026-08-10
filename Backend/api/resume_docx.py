"""
Render the *tailored* resume to DOCX and TXT.

This exists because the two halves of the build used to disagree. `resume_build.py`
rebuilds whatever file you uploaded; `latex_resume.py` scores, reorders and
rewrites your curated profile against the posting. Downloading "PDF" therefore
gave you the plain rebuild of your upload while ".TEX" gave you the tailored
document — two different resumes from one button row.

Everything here takes the same `ResumeContent` the LaTeX template renders, so
the PDF, the DOCX and the TXT are the same document in three containers.

The layout mirrors the reference resume exactly, as the LaTeX template does:
17pt bold name, the contact line with full URLs, bold ALL-CAPS headings on a
hairline rule, italic dates flush right on a native tab stop, organisation and
location together on one italic line, Technical Skills as bold-label lines,
and one Achievements & Certifications section that also carries languages.
Dates sit on a tab stop rather than in a table, because parsers read table
cells out of order.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from . import resume_style

# 8.5in page, 0.5in margins each side.
TEXT_WIDTH_IN = 7.5
BODY_PT = 10.5
NAME_PT = 17
HEAD_PT = 12
SMALL_PT = 9.5


# Both delegate to resume_style so this renderer and the LaTeX one share one
# definition of "how skills split" and "what order the contact line goes in".
def _rows(skills: list[str]) -> list[tuple[str, str]]:
    return [(r["label"], r["value"]) for r in resume_style.skill_rows(skills)]


def _contact_line(content: Any) -> str:
    return " | ".join(resume_style.contact_parts(content))


def _achievements(content: Any) -> list[tuple[str, str]]:
    """(bold label, text) bullets: awards, then certifications, then languages."""
    E = resume_style.enforce
    out: list[tuple[str, str]] = [("", E(a)) for a in content.awards or [] if a]
    certs = [E(c) for c in content.certifications or [] if c]
    if certs:
        out.append(("Certifications:", ", ".join(certs)))
    langs = [E(l) for l in content.languages or [] if l]
    if langs:
        out.append(("Languages:", ", ".join(langs)))
    return out


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def _hairline(paragraph) -> None:
    """A 0.5pt black rule under a heading — the DOCX twin of \\rule in LaTeX."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")        # eighths of a point
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def render_docx(content: Any, path: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    E = resume_style.enforce
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.04

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)

    def para(space_before: float = 0, space_after: float = 2, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            p.alignment = align
        return p

    def run(p, text: str, *, bold=False, italic=False, size=BODY_PT):
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
        return r

    def heading(text: str) -> None:
        p = para(space_before=9, space_after=3)
        run(p, text.upper(), bold=True, size=HEAD_PT)
        _hairline(p)

    def bullet(text: str, *, label: str = "") -> None:
        p = para(space_after=1.5)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        run(p, "•  ")
        if label:
            run(p, label + " ", bold=True)
        run(p, text)

    def dated(left: str, right: str, *, bold_left=False) -> None:
        """Bold left text with the dates flush right in italics — the
        reference's entry line, on a native right-aligned tab stop."""
        p = para(space_after=0.5)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(TEXT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
        run(p, E(left), bold=bold_left)
        if right:
            run(p, "\t" + E(right), italic=True)

    # ---- header: name, then the full-URL contact line -------------------
    p = para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, content.name or "", bold=True, size=NAME_PT)
    line = _contact_line(content)
    if line:
        p = para(space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        run(p, line, size=SMALL_PT)

    # ---- body, in the reference section order ----------------------------
    if content.summary:
        heading("Summary")
        run(para(), E(content.summary))

    if content.experience:
        heading("Professional Experience")
        for block in content.experience:
            dated(block.role or "", block.period or "", bold_left=True)
            if block.company or block.location:
                where = ", ".join(x for x in (block.company, block.location) if x)
                run(para(space_after=1), E(where), italic=True)
            for b in block.bullets:
                bullet(E(b.text))

    if content.projects:
        heading("Projects")
        for block in content.projects:
            p = para(space_before=3, space_after=1)
            run(p, E(block.name or ""), bold=True)
            if block.tech:
                run(p, " | ", bold=True)
                run(p, E(block.tech), bold=True, italic=True)
            for b in block.bullets:
                bullet(E(b.text))

    rows = _rows(content.skills or [])
    if rows:
        heading("Technical Skills")
        for label, items in rows:
            p = para(space_after=1)
            run(p, f"{label}: ", bold=True)
            run(p, items)

    entries = getattr(content, "education_entries", None) or []
    if entries or content.education:
        heading("Education")
        for e in entries:
            place = ", ".join(x for x in (e.get("institution"), e.get("location")) if x)
            dated(place, e.get("period") or "", bold_left=True)
            if e.get("degree"):
                run(para(space_after=1), E(e["degree"]), italic=True)
            if e.get("courses"):
                run(para(space_after=1), "Relevant Courses: " + E(e["courses"]))
        for entry in content.education:
            run(para(), E(entry))

    ach = _achievements(content)
    if ach:
        heading("Achievements & Certifications")
        for label, text in ach:
            bullet(text, label=label)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


# --------------------------------------------------------------------------
# TXT
# --------------------------------------------------------------------------

def render_txt(content: Any) -> str:
    """
    Plain text of the same document.

    Underlines are '=' and bullets are '•': the house style bans the hyphen,
    and a row of them is exactly the kind of token an ATS indexes as a non-word.
    """
    E = resume_style.enforce
    out: list[str] = [content.name or ""]
    line = _contact_line(content)
    if line:
        out.append(line)
    out.append("")

    def head(title: str) -> None:
        out.append(title.upper())
        out.append("=" * len(title))

    if content.summary:
        head("Summary")
        out.append(E(content.summary))
        out.append("")

    if content.experience:
        head("Professional Experience")
        for block in content.experience:
            role = E(block.role or "")
            out.append(f"{role}   {E(block.period)}" if block.period else role)
            where = ", ".join(x for x in (block.company, block.location) if x)
            if where:
                out.append(E(where))
            out.extend(f"• {E(b.text)}" for b in block.bullets)
            out.append("")
        if out and out[-1] == "":
            out.pop()
        out.append("")

    if content.projects:
        head("Projects")
        for block in content.projects:
            name = E(block.name or "")
            out.append(f"{name} | {E(block.tech)}" if block.tech else name)
            out.extend(f"• {E(b.text)}" for b in block.bullets)
            out.append("")
        if out and out[-1] == "":
            out.pop()
        out.append("")

    rows = _rows(content.skills or [])
    if rows:
        head("Technical Skills")
        out.extend(f"{label}: {items}" for label, items in rows)
        out.append("")

    entries = getattr(content, "education_entries", None) or []
    if entries or content.education:
        head("Education")
        for e in entries:
            place = ", ".join(x for x in (e.get("institution"), e.get("location")) if x)
            out.append(f"{place}   {E(e.get('period') or '')}".rstrip())
            if e.get("degree"):
                out.append(E(e["degree"]))
            if e.get("courses"):
                out.append("Relevant Courses: " + E(e["courses"]))
        out.extend(E(x) for x in content.education)
        out.append("")

    ach = _achievements(content)
    if ach:
        head("Achievements & Certifications")
        out.extend(f"• {label} {text}" if label else f"• {text}" for label, text in ach)
        out.append("")
    return "\n".join(out).rstrip() + "\n"
