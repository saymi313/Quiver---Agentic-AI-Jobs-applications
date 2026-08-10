"""
Extract text, sections and layout warnings from an uploaded resume.

Supports PDF (pdfplumber, PyMuPDF fallback), DOCX (python-docx) and plain text.
The goal is not perfect fidelity but an ATS-shaped view of the document: what a
resume parser would actually be able to read.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

BULLET_CHARS = "•·▪●○–—-*‣◦➢»"
BULLET_RE = re.compile(rf"^\s*[{re.escape(BULLET_CHARS)}]\s+")

# The one email regex, shared with agent/sources.py. The {2,} tail rejects
# truncated artifacts like "a@b.c." that page scrapes produce.
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]{2,}")
PHONE_RE = re.compile(r"(?:\+?\d[\d\s().-]{7,}\d)")
URL_RE = re.compile(r"(?:https?://|www\.)[^\s,;)]+", re.I)
LINKEDIN_RE = re.compile(r"linkedin\.com/[^\s,;)]+", re.I)
GITHUB_RE = re.compile(r"github\.com/[^\s,;)]+", re.I)

DATE_RANGE_RE = re.compile(
    r"("
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"|\d{1,2}\s*/\s*\d{1,2}\s*/\s*\d{4}"
    r"|\d{1,2}\s*/\s*\d{4}"
    r"|\b(?:19|20)\d{2}\b"
    r")"
    r"\s*(?:-|–|—|to|until|through)\s*"
    r"("
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"|\d{1,2}\s*/\s*\d{1,2}\s*/\s*\d{4}"
    r"|\d{1,2}\s*/\s*\d{4}"
    r"|\b(?:19|20)\d{2}\b"
    r"|present|current|now|ongoing|to date"
    r")",
    re.I,
)

# Canonical section -> heading spellings we accept.
SECTION_ALIASES: dict[str, list[str]] = {
    "summary": [
        "summary", "professional summary", "career summary", "profile",
        "professional profile", "about me", "about", "objective",
        "career objective", "personal statement", "overview", "executive summary",
    ],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment", "employment history", "work history", "career history",
        "relevant experience", "professional background", "engineering experience",
    ],
    "education": [
        "education", "academic background", "academics", "qualifications",
        "education and training", "academic qualifications", "educational background",
    ],
    "skills": [
        "skills", "technical skills", "core skills", "core competencies",
        "competencies", "key skills", "skills and abilities", "technologies",
        "tech stack", "technical proficiencies", "areas of expertise", "expertise",
    ],
    "projects": [
        "projects", "personal projects", "selected projects", "key projects",
        "side projects", "portfolio", "notable projects", "academic projects",
    ],
    "certifications": [
        "certifications", "certificates", "licenses", "licenses and certifications",
        "professional certifications", "courses", "training", "coursework",
    ],
    "awards": [
        "awards", "honors", "achievements", "accomplishments", "recognition",
        "honors and awards",
    ],
    "publications": ["publications", "papers", "research", "patents"],
    "volunteer": ["volunteer", "volunteering", "community", "extracurricular", "activities"],
    "languages": ["languages", "language proficiency"],
    "interests": ["interests", "hobbies", "personal interests"],
    "references": ["references", "referees"],
}

HEADING_LOOKUP: dict[str, str] = {}
for _canon, _names in SECTION_ALIASES.items():
    for _n in _names:
        HEADING_LOOKUP[_n] = _canon

# Heading spellings that parse reliably everywhere — anything outside this set
# is flagged as unusual rather than wrong.
SAFE_HEADINGS = {
    "summary": {"summary", "professional summary", "profile", "objective"},
    "experience": {"experience", "work experience", "professional experience",
                   "employment history", "work history"},
    "education": {"education"},
    "skills": {"skills", "technical skills", "core skills", "core competencies"},
    "projects": {"projects", "key projects", "selected projects"},
    "certifications": {"certifications", "certificates"},
    "awards": {"awards", "achievements", "honors and awards", "awards and certifications",
               "achievements and certifications", "certifications and awards"},
    "publications": {"publications"},
    "volunteer": {"volunteer", "volunteer experience"},
    "languages": {"languages"},
    "interests": {"interests"},
    "references": {"references"},
}

LINK_LABEL_RE = re.compile(
    r"\b(linkedin|github|gitlab|portfolio|website|behance|dribbble|medium|twitter|x\.com|"
    r"stack\s?overflow|kaggle|resume|cv)\b", re.I
)

# Headings an ATS is guaranteed to understand.
STANDARD_HEADINGS = {
    "summary": "PROFESSIONAL SUMMARY",
    "skills": "SKILLS",
    "experience": "PROFESSIONAL EXPERIENCE",
    "projects": "PROJECTS",
    "education": "EDUCATION",
    "certifications": "CERTIFICATIONS",
    "awards": "AWARDS",
    "publications": "PUBLICATIONS",
    "volunteer": "VOLUNTEER EXPERIENCE",
    "languages": "LANGUAGES",
    "interests": "INTERESTS",
    "references": "REFERENCES",
}

SECTION_ORDER = [
    "summary", "skills", "experience", "projects", "education",
    "certifications", "awards", "publications", "volunteer",
    "languages", "interests", "references",
]


@dataclass
class ExperienceEntry:
    header: str = ""
    title: str = ""
    organization: str = ""
    period: str = ""
    location: str = ""
    bullets: list[str] = field(default_factory=list)


@dataclass
class ParsedResume:
    raw_text: str = ""
    lines: list[str] = field(default_factory=list)
    source: str = ""
    page_count: int = 0
    word_count: int = 0
    contact: dict[str, str] = field(default_factory=dict)
    name: str = ""
    headline: str = ""
    sections: dict[str, list[str]] = field(default_factory=dict)
    heading_map: dict[str, str] = field(default_factory=dict)
    experience: list[ExperienceEntry] = field(default_factory=list)
    projects: list[ExperienceEntry] = field(default_factory=list)
    layout_warnings: list[dict[str, str]] = field(default_factory=list)
    # URLs that exist only as clickable annotations, not as visible characters
    embedded_links: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "source": self.source,
            "pageCount": self.page_count,
            "wordCount": self.word_count,
            "name": self.name,
            "headline": self.headline,
            "contact": self.contact,
            "sectionsFound": [s for s in SECTION_ORDER if self.sections.get(s)],
            "headingMap": self.heading_map,
            "embeddedLinks": self.embedded_links,
            "experience": [
                {
                    "header": e.header,
                    "title": e.title,
                    "organization": e.organization,
                    "period": e.period,
                    "bullets": e.bullets,
                }
                for e in self.experience
            ],
            "projects": [
                {"header": p.header, "title": p.title, "bullets": p.bullets}
                for p in self.projects
            ],
            "layoutWarnings": self.layout_warnings,
        }


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------

def _text_quality(text: str) -> float:
    """Lower is better. Penalises 'gluedtogetherwords' from bad space recovery."""
    tokens = [t for t in re.findall(r"[A-Za-z][A-Za-z'-]*", text) if t]
    if len(tokens) < 20:
        return 1.0
    glued = sum(1 for t in tokens if len(t) > 17)
    return glued / len(tokens)


def _extract_pdf(path: Path) -> tuple[str, int, list[dict[str, str]], list[str]]:
    warnings: list[dict[str, str]] = []
    links: list[str] = []
    text = ""
    pages = 0

    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None

    if pdfplumber is not None:
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages = len(pdf.pages)
                chunks: list[str] = []
                table_pages = 0
                image_pages = 0
                column_pages = 0
                for page in pdf.pages:
                    chunks.append(page.extract_text() or "")
                    try:
                        for h in (page.hyperlinks or []):
                            uri = h.get("uri")
                            if uri and uri not in links:
                                links.append(uri)
                        for a in (page.annots or []):
                            uri = a.get("uri")
                            if uri and uri not in links:
                                links.append(uri)
                    except Exception:
                        pass
                    try:
                        if page.find_tables():
                            table_pages += 1
                    except Exception:
                        pass
                    if getattr(page, "images", None):
                        image_pages += 1
                    if _looks_multi_column(page):
                        column_pages += 1
                text = "\n".join(chunks)

                if table_pages:
                    warnings.append({
                        "id": "tables",
                        "severity": "high",
                        "title": f"Tables detected on {table_pages} page(s)",
                        "detail": "Many ATS parsers read table cells out of order or drop them entirely. "
                                  "Content should sit in a single linear column.",
                    })
                if column_pages:
                    warnings.append({
                        "id": "columns",
                        "severity": "high",
                        "title": f"Multi-column layout on {column_pages} page(s)",
                        "detail": "Side-by-side columns get interleaved during parsing, which scrambles "
                                  "job titles, dates and bullets.",
                    })
                if image_pages:
                    warnings.append({
                        "id": "images",
                        "severity": "medium",
                        "title": f"Images / graphics on {image_pages} page(s)",
                        "detail": "Logos, photos, icons and skill-rating bars carry no text. Anything "
                                  "that matters must also exist as real characters.",
                    })
        except Exception as exc:  # pragma: no cover - depends on file
            warnings.append({
                "id": "pdf_parse",
                "severity": "medium",
                "title": "PDF layout analysis failed",
                "detail": f"Fell back to plain text extraction ({exc}).",
            })

    # pdfplumber gives the best layout signals but sometimes drops word spacing on
    # justified text; PyMuPDF is usually cleaner. Take whichever reads better.
    try:
        import fitz  # PyMuPDF

        with fitz.open(str(path)) as doc:
            pages = pages or doc.page_count
            alt = "\n".join(page.get_text("text") for page in doc)
            for page in doc:
                for link in page.get_links():
                    uri = link.get("uri")
                    if uri and uri not in links:
                        links.append(uri)
        if alt.strip() and (len(text.strip()) < 80 or _text_quality(alt) < _text_quality(text) - 0.005):
            text = alt
    except Exception:
        pass

    if len(text.strip()) < 80:
        warnings.append({
            "id": "no_text_layer",
            "severity": "critical",
            "title": "Almost no machine-readable text",
            "detail": "This PDF is likely a scan or exported as an image. An ATS will read it as a "
                      "blank document. Export from the original editor as a text-based PDF.",
        })

    return text, pages, warnings, links


def _looks_multi_column(page) -> bool:
    """Heuristic: two dense word clusters separated by a vertical gutter."""
    try:
        words = page.extract_words()
    except Exception:
        return False
    if len(words) < 40:
        return False
    width = float(page.width or 0)
    if width <= 0:
        return False

    mids = sorted(((float(w["x0"]) + float(w["x1"])) / 2.0) / width for w in words)
    left = [m for m in mids if m < 0.45]
    right = [m for m in mids if m > 0.55]
    if len(left) < len(mids) * 0.25 or len(right) < len(mids) * 0.25:
        return False
    # A real gutter means very few words straddle the middle band.
    middle = [m for m in mids if 0.45 <= m <= 0.55]
    return len(middle) < len(mids) * 0.04


def _extract_docx(path: Path) -> tuple[str, int, list[dict[str, str]], list[str]]:
    from docx import Document

    warnings: list[dict[str, str]] = []
    links: list[str] = []
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]

    try:
        for rel in doc.part.rels.values():
            if rel.reltype.endswith("/hyperlink") and rel.is_external:
                target = str(rel.target_ref)
                if target and target not in links:
                    links.append(target)
    except Exception:
        pass

    if doc.tables:
        warnings.append({
            "id": "tables",
            "severity": "high",
            "title": f"{len(doc.tables)} table(s) in the document",
            "detail": "Tables are the most common cause of scrambled ATS output. Move the content "
                      "into ordinary paragraphs.",
        })
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("  ".join(cells))

    for section in doc.sections:
        head = "\n".join(p.text for p in section.header.paragraphs).strip()
        foot = "\n".join(p.text for p in section.footer.paragraphs).strip()
        if head or foot:
            warnings.append({
                "id": "header_footer",
                "severity": "medium",
                "title": "Content in the page header/footer",
                "detail": "Several parsers ignore headers and footers outright. Contact details in "
                          "particular must live in the document body.",
            })
            parts.extend([head, foot])
            break

    try:
        if doc.inline_shapes:
            warnings.append({
                "id": "images",
                "severity": "medium",
                "title": f"{len(doc.inline_shapes)} inline image(s)",
                "detail": "Images are invisible to keyword matching.",
            })
    except Exception:
        pass

    return "\n".join(parts), 0, warnings, links


def extract_text(path: Path) -> tuple[str, int, str, list[dict[str, str]], list[str]]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        text, pages, warns, links = _extract_pdf(path)
        return text, pages, "pdf", warns, links
    if ext == ".docx":
        text, pages, warns, links = _extract_docx(path)
        return text, pages, "docx", warns, links
    if ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace"), 0, ext.lstrip("."), [], []
    raise ValueError(f"Unsupported resume format: {ext or '(none)'}")


# --------------------------------------------------------------------------
# Structure
# --------------------------------------------------------------------------

# Fonts that ship no ToUnicode map make extractors emit the raw glyph id as
# literal text: "(cid:127) Built the thing". 127/129/149/183 are the bullet
# slots in the usual symbol fonts; anything else is a glyph we cannot recover,
# so it becomes a space rather than garbage in the middle of a word.
CID_RE = re.compile(r"\(cid:(\d+)\)")
CID_BULLETS = {127, 129, 149, 183, 8226, 61623, 61590}


def _normalize_cid(line: str) -> str:
    if "(cid:" not in line:
        return line

    def repl(m: re.Match) -> str:
        return "• " if int(m.group(1)) in CID_BULLETS else " "

    line = CID_RE.sub(repl, line)
    return re.sub(r"\s{2,}", " ", line).strip()


def _clean_lines(text: str) -> list[str]:
    out: list[str] = []
    for raw in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw.replace("\t", " ").replace("\xa0", " ")
        line = _normalize_cid(line)
        line = re.sub(r"[ ]{2,}", "  ", line).rstrip()
        out.append(line)
    # collapse runs of blank lines
    collapsed: list[str] = []
    blank = False
    for line in out:
        if line.strip():
            collapsed.append(line.strip())
            blank = False
        elif not blank:
            collapsed.append("")
            blank = True
    while collapsed and not collapsed[0]:
        collapsed.pop(0)
    while collapsed and not collapsed[-1]:
        collapsed.pop()
    return _rejoin_hyphenation(collapsed)


def _rejoin_hyphenation(lines: list[str]) -> list[str]:
    """PDF line wrapping splits words as 'de-' / 'livering'. Put them back."""
    out: list[str] = []
    for line in lines:
        if (out and out[-1].endswith("-") and len(out[-1]) > 2
                and out[-1][-2].isalpha() and line[:1].islower()
                and not _normalize_heading(out[-1])):
            out[-1] = out[-1][:-1] + line
        else:
            out.append(line)
    return out


def _normalize_heading(line: str) -> str | None:
    """Return the canonical section name if this line is a section heading."""
    s = line.strip()
    if not s or len(s) > 60:
        return None
    if BULLET_RE.match(s):
        return None
    core = re.sub(r"^[^A-Za-z]+", "", s)
    core = re.sub(r"[^A-Za-z& ]+$", "", core).strip()
    if not core:
        return None
    key = re.sub(r"\s+", " ", core.lower()).strip(" :-–—")
    if len(s.split()) > 7:
        return None
    if key in HEADING_LOOKUP:
        return HEADING_LOOKUP[key]
    normalized = key.replace(" & ", " and ")
    if normalized in HEADING_LOOKUP:
        return HEADING_LOOKUP[normalized]
    # Compound headings such as "Achievements & Certifications" or
    # "Education / Training" — anchor on the first half we recognise.
    for part in re.split(r"\s*(?:&|/|,|\band\b)\s*", key):
        part = part.strip()
        if part in HEADING_LOOKUP:
            return HEADING_LOOKUP[part]
    return None


def _split_sections(lines: list[str]) -> tuple[list[str], dict[str, list[str]], dict[str, str]]:
    header_block: list[str] = []
    sections: dict[str, list[str]] = {}
    heading_map: dict[str, str] = {}
    current: str | None = None

    for line in lines:
        canon = _normalize_heading(line)
        if canon:
            current = canon
            sections.setdefault(canon, [])
            heading_map.setdefault(canon, line.strip())
            continue
        if current is None:
            header_block.append(line)
        else:
            sections[current].append(line)

    for key in list(sections):
        body = sections[key]
        while body and not body[0]:
            body.pop(0)
        while body and not body[-1]:
            body.pop()
        if not body:
            sections.pop(key)
    return header_block, sections, heading_map


def _extract_contact(text: str, header_block: list[str]) -> tuple[dict[str, str], str, str]:
    head_text = "\n".join(header_block[:12]) or text[:1200]
    contact: dict[str, str] = {}

    m = EMAIL_RE.search(head_text) or EMAIL_RE.search(text)
    if m:
        contact["email"] = m.group(0).strip(".,;")

    for cand in PHONE_RE.finditer(head_text) or []:
        digits = re.sub(r"\D", "", cand.group(0))
        if 9 <= len(digits) <= 15:
            contact["phone"] = cand.group(0).strip()
            break
    if "phone" not in contact:
        for cand in PHONE_RE.finditer(text):
            digits = re.sub(r"\D", "", cand.group(0))
            if 9 <= len(digits) <= 15:
                contact["phone"] = cand.group(0).strip()
                break

    m = LINKEDIN_RE.search(text)
    if m:
        contact["linkedin"] = m.group(0).rstrip(".,;)")
    m = GITHUB_RE.search(text)
    if m:
        contact["github"] = m.group(0).rstrip(".,;)")

    others = [
        u.rstrip(".,;)") for u in URL_RE.findall(text)
        if "linkedin.com" not in u.lower() and "github.com" not in u.lower()
    ]
    if others:
        contact["website"] = others[0]

    loc = _guess_location(head_text)
    if loc:
        contact["location"] = loc

    name = ""
    headline = ""
    for line in header_block[:6]:
        s = line.strip()
        if not s or EMAIL_RE.search(s) or URL_RE.search(s) or LINK_LABEL_RE.search(s):
            continue
        words = s.split()
        if not name and 1 < len(words) <= 5 and not re.search(r"\d", s):
            letters = re.sub(r"[^A-Za-z ]", "", s)
            if letters and (s.isupper() or s == s.title() or letters.istitle()):
                name = re.sub(r"\s+", " ", s).strip()
                continue
        if name and not headline and 2 <= len(words) <= 12 and not re.search(r"\d{4}", s):
            headline = s
            break
    return contact, name, headline


def _guess_location(text: str) -> str:
    for line in text.split("\n"):
        # Contact rows are usually "City, Country | phone | email" — check each cell.
        for cell in re.split(r"\||•|·|;", line):
            s = cell.strip()
            if not s or EMAIL_RE.search(s) or URL_RE.search(s) or len(s) > 60:
                continue
            m = re.search(
                r"\b([A-Z][a-zA-Z.\-]+(?:\s[A-Z][a-zA-Z.\-]+)*,\s*[A-Z][a-zA-Z.\-]+"
                r"(?:\s[A-Z][a-zA-Z.\-]+)*)\b", s)
            if m and len(m.group(1)) < 60 and not DATE_RANGE_RE.search(m.group(1)):
                return m.group(1)
    return ""


def _looks_like_entry_header(line: str) -> bool:
    if BULLET_RE.match(line):
        return False
    if DATE_RANGE_RE.search(line):
        return True
    if re.search(r"\b(19|20)\d{2}\b", line) and len(line.split()) <= 14:
        return True
    if "|" in line and len(line.split()) <= 16:
        return True
    if re.search(r"\bat\b", line, re.I) and len(line.split()) <= 12 and line[:1].isupper():
        return True
    return False


def _parse_entries(body: list[str]) -> list[ExperienceEntry]:
    """
    Group a section body into role/project entries.

    PDF extraction often emits a role header as several separate short lines
    (title, dates, employer, location), so short lines seen before a date line
    are held in `pending` and folded into the header rather than dropped.
    """
    entries: list[ExperienceEntry] = []
    current: ExperienceEntry | None = None
    pending: list[str] = []

    def start(header_parts: list[str]) -> ExperienceEntry:
        nonlocal current
        current = ExperienceEntry(header=" | ".join(p for p in header_parts if p))
        entries.append(current)
        return current

    for line in body:
        if not line:
            continue

        if BULLET_RE.match(line):
            text = BULLET_RE.sub("", line).strip()
            if not text:
                continue
            if current is None or (pending and current.bullets):
                start(pending[-3:])
                pending = []
            current.bullets.append(text)
            continue

        if _looks_like_entry_header(line):
            start(pending[-3:] + [line.strip()])
            pending = []
            continue

        if current is not None and current.bullets:
            if line[:1].islower():
                current.bullets[-1] = f"{current.bullets[-1]} {line.strip()}".strip()
            elif len(line.split()) > 12:
                current.bullets.append(line.strip())
            else:
                pending.append(line.strip())
            continue

        if current is not None and not current.bullets:
            current.header = f"{current.header} | {line.strip()}".strip(" |")
            continue

        pending.append(line.strip())

    if pending and not entries:
        entries.append(ExperienceEntry(header=" | ".join(pending)))

    entries = _merge_orphan_headers(entries)
    for e in entries:
        _split_header(e)
    return [e for e in entries if e.header or e.bullets]


def _merge_orphan_headers(entries: list[ExperienceEntry]) -> list[ExperienceEntry]:
    """
    Template resumes (Europass and friends) split one role across two header
    blocks — dates in one cell, employer in the next. Fold a bullet-less entry
    into the entry that follows it.
    """
    merged: list[ExperienceEntry] = []
    i = 0
    while i < len(entries):
        entry = entries[i]
        if not entry.bullets and i + 1 < len(entries) and entries[i + 1].bullets:
            nxt = entries[i + 1]
            nxt.header = f"{entry.header} | {nxt.header}".strip(" |")
            i += 1
            continue
        merged.append(entry)
        i += 1
    return merged


def _split_header(entry: ExperienceEntry) -> None:
    header = entry.header
    m = DATE_RANGE_RE.search(header)
    if m:
        entry.period = m.group(0).strip()
        header = (header[: m.start()] + " " + header[m.end():]).strip(" ,|–—-")
    parts = [p.strip(" ,|–—-") for p in re.split(r"\||—|–|•", header) if p.strip(" ,|–—-")]
    if len(parts) >= 2:
        entry.title = parts[0]
        entry.organization = parts[1]
        if len(parts) >= 3:
            entry.location = parts[2]
    elif parts:
        chunk = parts[0]
        m2 = re.split(r"\bat\b|,", chunk, maxsplit=1, flags=re.I)
        if len(m2) == 2:
            entry.title, entry.organization = m2[0].strip(), m2[1].strip()
        else:
            entry.title = chunk
    entry.header = " | ".join(x for x in [entry.title, entry.organization, entry.location, entry.period] if x) \
        or entry.header


def parse_resume(path: Path) -> ParsedResume:
    text, pages, source, warnings, links = extract_text(path)
    lines = _clean_lines(text)
    header_block, sections, heading_map = _split_sections(lines)
    contact, name, headline = _extract_contact(text, header_block)

    parsed = ParsedResume(
        raw_text=text,
        lines=lines,
        source=source,
        page_count=pages,
        word_count=len(re.findall(r"[A-Za-z0-9][A-Za-z0-9'./+-]*", text)),
        contact=contact,
        name=name,
        headline=headline,
        sections=sections,
        heading_map=heading_map,
        layout_warnings=warnings,
        embedded_links=[l for l in links if not l.lower().startswith("mailto:")],
    )
    parsed.experience = _parse_entries(sections.get("experience", []))
    parsed.projects = _parse_entries(sections.get("projects", []))

    # A profile URL that exists only as a clickable annotation is invisible to a
    # text-based parser — worth calling out, and worth restoring in the rebuild.
    # "Hidden" means the URL exists only as a clickable annotation. Compare on
    # the bare host+path so a link written as visible text "linkedin.com/in/me"
    # is not flagged just because the annotation carries the https://www. prefix.
    def _bare(u: str) -> str:
        return re.sub(r"^https?://(www\.)?", "", (u or "").strip()).rstrip("/").lower()

    body = _bare_text = text.lower()
    hidden = [l for l in parsed.embedded_links if _bare(l) and _bare(l) not in body]
    if hidden:
        parsed.layout_warnings.append({
            "id": "hyperlink_only",
            "severity": "medium",
            "title": f"{len(hidden)} profile link(s) exist only as clickable hyperlinks",
            "detail": "Words like \"LinkedIn\" or \"Portfolio\" hyperlinked to a URL look fine to a "
                      "human but give a text parser nothing. Write the URL out: "
                      "linkedin.com/in/your-handle.",
        })
        for url in hidden:
            low = url.lower()
            if "linkedin.com" in low:
                parsed.contact.setdefault("linkedinHidden", url)
            elif "github.com" in low:
                parsed.contact.setdefault("githubHidden", url)
            else:
                parsed.contact.setdefault("websiteHidden", url)

    if not sections:
        parsed.layout_warnings.append({
            "id": "no_sections",
            "severity": "high",
            "title": "No recognisable section headings",
            "detail": "Parsers anchor on headings like EXPERIENCE, EDUCATION and SKILLS. Without them "
                      "the whole document is treated as one undifferentiated blob.",
        })
    else:
        odd = [
            raw for canon, raw in heading_map.items()
            if re.sub(r"\s+", " ", raw.strip().lower().strip(" :-–—")) not in SAFE_HEADINGS.get(canon, set())
        ]
        if odd:
            parsed.layout_warnings.append({
                "id": "nonstandard_headings",
                "severity": "low",
                "title": "Non-standard section headings",
                "detail": "Recognised but unusual: " + ", ".join(sorted(set(odd))[:6]) +
                          ". Plain headings (SKILLS, PROFESSIONAL EXPERIENCE, EDUCATION) parse most reliably.",
            })

    if pages and pages > 2:
        parsed.layout_warnings.append({
            "id": "length",
            "severity": "low",
            "title": f"{pages} pages",
            "detail": "One page for under ~8 years of experience, two at most otherwise.",
        })
    return parsed
