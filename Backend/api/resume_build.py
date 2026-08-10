"""
Rebuild a parsed resume into an ATS-safe document and render it to PDF, DOCX
and TXT.

Rules that drive every decision here:
  * one linear column, no tables, no text boxes, no headers/footers, no images
  * standard section headings an ATS is guaranteed to recognise
  * contact details as plain text in the body
  * content is reordered and relabelled, never invented — every bullet that
    comes out was already in the uploaded resume (or was rewritten by the
    optional LLM pass, which is constrained to the same facts)
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from . import resume_style
from .ats import ALIAS_TO_CANON, SKILL_LEXICON, _alias_pattern, _norm
from .resume_parse import STANDARD_HEADINGS, ExperienceEntry, ParsedResume

SECTION_RENDER_ORDER = [
    "summary", "skills", "experience", "projects", "education",
    "certifications", "awards", "publications", "volunteer", "languages",
]


@dataclass
class BuiltResume:
    name: str = ""
    headline: str = ""
    contact_lines: list[str] = field(default_factory=list)
    summary: str = ""
    skills: list[str] = field(default_factory=list)
    experience: list[dict[str, Any]] = field(default_factory=list)
    projects: list[dict[str, Any]] = field(default_factory=list)
    education: list[str] = field(default_factory=list)
    extra: dict[str, list[str]] = field(default_factory=dict)
    notes: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "name": self.name,
            "headline": self.headline,
            "contactLines": self.contact_lines,
            "summary": self.summary,
            "skills": self.skills,
            "experience": self.experience,
            "projects": self.projects,
            "education": self.education,
            "extra": self.extra,
            "notes": self.notes,
        }


# --------------------------------------------------------------------------
# Assembly
# --------------------------------------------------------------------------

def _bullet_relevance(bullet: str, jd: dict[str, Any]) -> float:
    norm = _norm(bullet)
    score = 0.0
    for kw in jd["keywords"]:
        for alias in kw.get("aliases", [kw["term"]]):
            if _alias_pattern(alias).search(norm):
                score += kw["weight"]
                break
    if re.search(r"\d", bullet):
        score += 0.8
    return score


def _contact_lines(parsed: ParsedResume) -> list[str]:
    c = parsed.contact
    primary = [c.get("location", ""), c.get("phone", ""), c.get("email", "")]
    primary = [p for p in primary if p]
    # Fall back to hyperlink targets so links that were invisible to a parser
    # come back as real, readable text.
    links = [
        c.get("linkedin") or c.get("linkedinHidden", ""),
        c.get("github") or c.get("githubHidden", ""),
        c.get("website") or c.get("websiteHidden", ""),
    ]
    links = [re.sub(r"^https?://(www\.)?", "", l).rstrip("/") for l in links if l]
    out = []
    if primary:
        out.append(" | ".join(primary))
    if links:
        out.append(" | ".join(links))
    return out


def _skills_lines(parsed: ParsedResume, jd: dict[str, Any], match: dict[str, Any]) -> list[str]:
    """Existing skills, plus a JD-aligned core line built ONLY from matched terms."""
    lines: list[str] = []

    core = [m["term"] for m in sorted(match["matched"], key=lambda k: -k["weight"])
            if m["category"] in ("hard", "soft")][:16]
    if core:
        lines.append("Job Matched Skills: " + ", ".join(_pretty(t) for t in core))

    existing = [l.strip() for l in parsed.sections.get("skills", []) if l.strip()]
    seen_norm = {_norm(l) for l in lines}
    for line in existing:
        cleaned = re.sub(r"^[•·▪●○–—*\-]\s*", "", line).strip()
        if not cleaned or _norm(cleaned) in seen_norm:
            continue
        lines.append(cleaned)
        seen_norm.add(_norm(cleaned))

    if len(lines) == 1 and not existing:
        # No SKILLS section existed — derive one from what the resume actually proves.
        resume_hits = _resume_skill_terms(parsed.raw_text)
        rest = [t for t in resume_hits if t not in core][:20]
        if rest:
            lines.append("Additional: " + ", ".join(_pretty(t) for t in rest))
    return lines


def _resume_skill_terms(text: str) -> list[str]:
    norm = _norm(text)
    found: dict[str, int] = {}
    for alias, canon in ALIAS_TO_CANON.items():
        if _alias_pattern(alias).search(norm):
            found[canon] = found.get(canon, 0) + 1
    return sorted(found, key=lambda t: (-found[t], t))


_PRETTY = {
    "ci/cd": "CI/CD", "aws": "AWS", "gcp": "GCP", "sql": "SQL", "nlp": "NLP",
    "llm": "LLMs", "api design": "API Design", "rest api": "REST APIs",
    "ux": "UX", "ui design": "UI Design", "seo": "SEO", "crm": "CRM", "erp": "ERP",
    "cms": "CMS", "saas": "SaaS", "etl": "ETL", "tdd": "TDD", "mlops": "MLOps",
    "html": "HTML", "css": "CSS", "php": "PHP", "grpc": "gRPC", "graphql": "GraphQL",
    "node.js": "Node.js", "next.js": "Next.js", ".net": ".NET", "c#": "C#", "c++": "C++",
    "power bi": "Power BI", "e-commerce": "Ecommerce", "real-time": "Real time",
    "scikit-learn": "scikit learn", "pytorch": "PyTorch", "tensorflow": "TensorFlow",
    "mongodb": "MongoDB", "postgresql": "PostgreSQL", "mysql": "MySQL", "nestjs": "NestJS",
    "fastapi": "FastAPI", "javascript": "JavaScript", "typescript": "TypeScript",
}


def _pretty(term: str) -> str:
    if term in _PRETTY:
        return _PRETTY[term]
    if len(term) <= 3 and term.isalpha():
        return term.upper()
    return " ".join(w if w in _PRETTY else w.capitalize() for w in term.split(" "))


def _entry_to_dict(e: ExperienceEntry, jd: dict[str, Any], reorder: bool) -> dict[str, Any]:
    bullets = list(e.bullets)
    if reorder and jd and len(bullets) > 1:
        ranked = sorted(
            enumerate(bullets),
            key=lambda pair: (-_bullet_relevance(pair[1], jd), pair[0]),
        )
        bullets = [b for _, b in ranked]
    return {
        "title": e.title or e.header,
        "organization": e.organization,
        "location": e.location,
        "period": e.period,
        "bullets": [re.sub(r"\s+", " ", b).strip() for b in bullets if b.strip()],
    }


def build_resume(
    parsed: ParsedResume,
    jd: dict[str, Any] | None,
    match: dict[str, Any] | None,
    *,
    reorder_bullets: bool = True,
    llm: dict[str, Any] | None = None,
) -> BuiltResume:
    built = BuiltResume()
    built.name = parsed.name or "Your Name"
    built.headline = parsed.headline
    built.contact_lines = _contact_lines(parsed)

    summary = " ".join(parsed.sections.get("summary", [])).strip()
    summary = re.sub(r"\s+", " ", summary)
    if llm and llm.get("summary"):
        summary = llm["summary"].strip()
        built.notes.append("Summary rewritten by the AI pass, constrained to facts already in your resume.")
    if not summary:
        top = [m["term"] for m in sorted((match or {}).get("matched", []), key=lambda k: -k["weight"])][:6]
        if top:
            role = (jd or {}).get("title") or built.headline or "engineer"
            summary = (f"{role} with hands-on delivery across "
                       + ", ".join(_pretty(t) for t in top) + ".")
            built.notes.append("No summary section existed — generated one from skills already "
                               "evidenced in your resume. Review the wording before sending.")
    built.summary = summary

    if match and jd:
        built.skills = _skills_lines(parsed, jd, match)
    else:
        built.skills = [l.strip() for l in parsed.sections.get("skills", []) if l.strip()]

    llm_bullets: dict[str, str] = {}
    if llm and llm.get("bullets"):
        for item in llm["bullets"]:
            original = (item.get("original") or "").strip()
            revised = (item.get("revised") or "").strip()
            if original and revised:
                llm_bullets[_norm(original)] = revised

    def _apply_llm(entry: dict[str, Any]) -> dict[str, Any]:
        if not llm_bullets:
            return entry
        entry["bullets"] = [llm_bullets.get(_norm(b), b) for b in entry["bullets"]]
        return entry

    built.experience = [
        _apply_llm(_entry_to_dict(e, jd or {}, reorder_bullets)) for e in parsed.experience
    ]
    built.projects = [
        _apply_llm(_entry_to_dict(p, jd or {}, reorder_bullets)) for p in parsed.projects
    ]
    if llm_bullets:
        built.notes.append(f"{len(llm_bullets)} bullet(s) rewritten by the AI pass — verify every "
                           "number before you send this.")

    built.education = [
        re.sub(r"^[•·▪●○–—*\-]\s*", "", l).strip()
        for l in parsed.sections.get("education", []) if l.strip()
    ]

    for key in ("certifications", "awards", "publications", "volunteer", "languages"):
        body = [re.sub(r"^[•·▪●○–—*\-]\s*", "", l).strip()
                for l in parsed.sections.get(key, []) if l.strip()]
        if body:
            built.extra[key] = body

    if reorder_bullets and any(len(e["bullets"]) > 1 for e in built.experience):
        built.notes.append("Bullets reordered inside each role so the most job-relevant work reads first.")
    built.notes.append("Rebuilt as a single-column document with standard headings, no tables, "
                       "no images and no header/footer content.")
    return built


# --------------------------------------------------------------------------
# Rendering
# --------------------------------------------------------------------------

def render_txt(built: BuiltResume) -> str:
    # Underlines are "=" and bullets are "•": the house style bans the hyphen
    # character outright, and a row of them is exactly the kind of token an ATS
    # indexes as a non-word.
    E = resume_style.enforce
    out: list[str] = [built.name]
    if built.headline:
        out.append(E(built.headline))
    out.extend(built.contact_lines)  # not enforced — URLs keep their hyphens
    out.append("")

    def section(title: str, body: list[str]) -> None:
        if not body:
            return
        out.append(title)
        out.append("=" * len(title))
        out.extend(E(b) for b in body)
        out.append("")

    if built.summary:
        section(STANDARD_HEADINGS["summary"], [built.summary])
    section(STANDARD_HEADINGS["skills"], built.skills)

    for label, entries in (("experience", built.experience), ("projects", built.projects)):
        if not entries:
            continue
        out.append(STANDARD_HEADINGS[label])
        out.append("=" * len(STANDARD_HEADINGS[label]))
        for e in entries:
            head = " | ".join(x for x in [e.get("title"), e.get("organization"), e.get("location")] if x)
            if e.get("period"):
                head = f"{head}   {e['period']}" if head else e["period"]
            out.append(E(head))
            for b in e["bullets"]:
                out.append(f"• {E(b)}")
            out.append("")
        if out and out[-1] == "":
            out.pop()
        out.append("")

    section(STANDARD_HEADINGS["education"], built.education)
    for key, body in built.extra.items():
        section(STANDARD_HEADINGS.get(key, key.upper()), body)
    return "\n".join(out).rstrip() + "\n"


def render_docx(built: BuiltResume, path: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.08

    for section in doc.sections:
        section.top_margin = Pt(39.6)
        section.bottom_margin = Pt(39.6)
        section.left_margin = Pt(39.6)
        section.right_margin = Pt(39.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(built.name)
    run.bold = True
    run.font.size = Pt(17)

    if built.headline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(resume_style.enforce(built.headline))
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Contact lines are not style-enforced: hyphens inside a URL or an email
    # address are part of the address, not house style.
    for line in built.contact_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(9.5)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)

    def body(text: str, bullet: bool = False) -> None:
        text = resume_style.enforce(text)
        p = doc.add_paragraph()
        if bullet:
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.first_line_indent = Pt(-10)
            p.add_run("•  " + text)
        else:
            p.add_run(text)

    if built.summary:
        heading(STANDARD_HEADINGS["summary"])
        body(built.summary)

    if built.skills:
        heading(STANDARD_HEADINGS["skills"])
        for line in built.skills:
            body(line)

    for key, entries in (("experience", built.experience), ("projects", built.projects)):
        if not entries:
            continue
        heading(STANDARD_HEADINGS[key])
        for e in entries:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            left = " | ".join(x for x in [e.get("title"), e.get("organization"), e.get("location")] if x)
            r = p.add_run(resume_style.enforce(left or e.get("title") or ""))
            r.bold = True
            if e.get("period"):
                r2 = p.add_run("   " + resume_style.enforce(e["period"]))
                r2.italic = True
                r2.font.size = Pt(9.5)
            for b in e["bullets"]:
                body(b, bullet=True)

    if built.education:
        heading(STANDARD_HEADINGS["education"])
        for line in built.education:
            body(line)

    for key, lines in built.extra.items():
        heading(STANDARD_HEADINGS.get(key, key.upper()))
        for line in lines:
            body(line)

    doc.save(str(path))
    return path


def render_pdf(built: BuiltResume, path: Path) -> Path:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    def esc(text: str) -> str:
        return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def styled(text: str) -> str:
        """House style, then XML escaping. Contact lines skip this — see below."""
        return esc(resume_style.enforce(text))

    # Times New Roman, pure black — the same house style the LaTeX template
    # enforces, so the two output paths cannot diverge. The name runs large:
    # the earlier 17pt was the single biggest reason this path read as plain.
    name_style = ParagraphStyle("name", fontName="Times-Bold", fontSize=20, leading=23,
                                alignment=TA_CENTER, textColor="#000000", spaceAfter=3)
    headline_style = ParagraphStyle("headline", fontName="Times-Roman", fontSize=11, leading=13,
                                    alignment=TA_CENTER, textColor="#000000", spaceAfter=2)
    contact_style = ParagraphStyle("contact", fontName="Times-Roman", fontSize=9.5, leading=12,
                                   alignment=TA_CENTER, textColor="#000000", spaceAfter=1)
    head_style = ParagraphStyle("head", fontName="Times-Bold", fontSize=12, leading=14,
                                spaceBefore=11, spaceAfter=1, textColor="#000000")
    body_style = ParagraphStyle("body", fontName="Times-Roman", fontSize=10.5, leading=13.5,
                                textColor="#000000", spaceAfter=2)
    entry_style = ParagraphStyle("entry", fontName="Times-Roman", fontSize=10.5, leading=13.5,
                                 textColor="#000000", spaceBefore=6, spaceAfter=1)
    # bulletFontName is a separate setting from fontName and defaults to
    # Helvetica, so without it the bullet glyphs alone break the Times rule.
    bullet_style = ParagraphStyle("bullet", parent=body_style, leftIndent=12, bulletIndent=2,
                                  bulletFontName="Times-Roman", bulletFontSize=10.5,
                                  spaceAfter=1.5)

    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        leftMargin=39.6, rightMargin=39.6, topMargin=39.6, bottomMargin=39.6,
        title=f"{built.name} - Resume", author=built.name,
    )

    flow: list[Any] = [Paragraph(esc(built.name), name_style)]
    if built.headline:
        flow.append(Paragraph(styled(built.headline), headline_style))
    for line in built.contact_lines:
        flow.append(Paragraph(esc(line), contact_style))

    def heading(text: str) -> None:
        # Heading over a full-width hairline, matching the LaTeX template.
        # Without the rule this path rendered as undifferentiated text blocks,
        # which is what "plain" looks like.
        flow.append(Paragraph(esc(text), head_style))
        flow.append(HRFlowable(width="100%", thickness=0.6, color="#000000",
                               spaceBefore=1, spaceAfter=4))

    if built.summary:
        heading(STANDARD_HEADINGS["summary"])
        flow.append(Paragraph(styled(built.summary), body_style))

    if built.skills:
        heading(STANDARD_HEADINGS["skills"])
        for line in built.skills:
            flow.append(Paragraph(styled(line), body_style))

    for key, entries in (("experience", built.experience), ("projects", built.projects)):
        if not entries:
            continue
        heading(STANDARD_HEADINGS[key])
        for e in entries:
            left = " | ".join(x for x in [e.get("title"), e.get("organization"), e.get("location")] if x)
            head = f"<b>{styled(left)}</b>"
            if e.get("period"):
                head += f" &nbsp;<i>{styled(e['period'])}</i>"
            flow.append(Paragraph(head, entry_style))
            for b in e["bullets"]:
                flow.append(Paragraph(styled(b), bullet_style, bulletText="•"))

    if built.education:
        heading(STANDARD_HEADINGS["education"])
        for line in built.education:
            flow.append(Paragraph(styled(line), body_style))

    for key, lines in built.extra.items():
        heading(STANDARD_HEADINGS.get(key, key.upper()))
        for line in lines:
            flow.append(Paragraph(styled(line), body_style))

    flow.append(Spacer(1, 2))
    doc.build(flow)
    return path
